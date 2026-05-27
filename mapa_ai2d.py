import streamlit as st
import pandas as pd
import requests
import xml.etree.ElementTree as ET
import geopandas as gpd
from shapely.geometry import Point 
import fiona
import folium
from streamlit_folium import st_folium
import os
import json
import time
from google import genai
from folium.plugins import Geocoder
import httpx 
import difflib  
import re       
from docx import Document
from io import BytesIO
from datetime import datetime

st.set_page_config(layout="wide", page_title="MPZP AI Master")

# --- OSTATECZNY WYTRYCH SSL (BRUTE FORCE) ---
oryginalny_client = httpx.Client.__init__
def pancerny_client(self, *args, **kwargs):
    kwargs['verify'] = False
    oryginalny_client(self, *args, **kwargs)
httpx.Client.__init__ = pancerny_client
# ---------------------------------------------

# --- NOWE FUNKCJE EKSPORTU ---
def generuj_word(tresc, nr_uchwaly, nazwa_planu):
    doc = Document()
    doc.add_heading('Raport z analizy MPZP (AI)', 0)
    
    # Metryczka
    p = doc.add_paragraph()
    p.add_run(f"Plan: {nazwa_planu}\n").bold = True
    p.add_run(f"Uchwała: {nr_uchwaly}\n")
    p.add_run(f"Data analizy: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    doc.add_heading('Wyniki analizy:', level=1)
    
    # Proste czyszczenie markdowna na potrzeby Worda
    czysty_tekst = str(tresc).replace('**', '').replace('* ', '• ')
    doc.add_paragraph(czysty_tekst)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def eksport_do_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Analiza MPZP')
        
        # Automatyczne ustawianie zawijania tekstu
        worksheet = writer.sheets['Analiza MPZP']
        for row in worksheet.iter_rows():
            for cell in row:
                if cell.column_letter == 'E': # Zakładając, że to kolumna z Raportem AI
                    cell.alignment = cell.alignment.copy(wrapText=True, vertical='top')
                else:
                    cell.alignment = cell.alignment.copy(vertical='top')
    return output.getvalue()

def eksport_do_csv(df):
    return df.to_csv(index=False).encode('utf-8')
    
def generuj_word_zbiorczy(df):
    doc = Document()
    doc.add_heading('Zbiorczy Raport z Analiz MPZP', 0)
    
    # Strona tytułowa i ręczny spis
    p = doc.add_paragraph()
    p.add_run(f"Data wygenerowania: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n").bold = True
    p.add_run(f"Liczba przeanalizowanych punktów: {len(df)}")
    
    doc.add_heading('Zestawienie analizowanych lokalizacji:', level=1)
    for idx, row in df.iterrows():
        doc.add_paragraph(f"{row['Lp.']}. {row['Name']} (Plan: {row['NAZWA PLANU']})", style='List Bullet')
        
    doc.add_page_break() # Przejście do nowej strony
    
    # Generowanie raportów dla każdego punktu
    for idx, row in df.iterrows():
        # Używamy Heading 1, aby Word mógł z tego łatwo zrobić automatyczny spis treści
        doc.add_heading(f"Punkt {row['Lp.']}: {row['Name']}", level=1)
        
        p_meta = doc.add_paragraph()
        p_meta.add_run(f"Plan: ").bold = True
        p_meta.add_run(f"{row['NAZWA PLANU']}\n")
        p_meta.add_run(f"Symbol: ").bold = True
        p_meta.add_run(f"{row['SYMBOL']}\n")
        
        doc.add_heading('Wynik analizy AI:', level=2)
        czysty_tekst = str(row['Raport AI']).replace('**', '').replace('* ', '• ')
        doc.add_paragraph(czysty_tekst)
        
        # Pusta linia odstępu między raportami
        doc.add_paragraph("")
        
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
# ---------------------------------------------

# --- 0. KONFIGURACJA AI I BAZY LOKALNEJ --- 
API_KEY = "..." # <--- Pamiętaj, by wkleić tu swój prawdziwy klucz przed uruchomieniem!
MODEL_ID = "gemini-2.5-flash"
PLIK_INDEKSU = "indeks_planow_hybrydowy.json"

try:
    client = genai.Client(api_key=API_KEY)
except Exception as e:
    st.error(f"❌ Błąd inicjalizacji klienta API Google! Szczegóły: {str(e)}")
    st.stop()

@st.cache_data
def wczytaj_indeks():
    try:
        with open(PLIK_INDEKSU, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return []

indeks_planow = wczytaj_indeks()

@st.cache_data
def wczytaj_granice_geojson():
    if os.path.exists("granice_mpzp_warszawa.geojson"):
        return gpd.read_file("granice_mpzp_warszawa.geojson")
    return None

def czytaj_tekst_planu(sciezka_folderu):
    if not sciezka_folderu or not os.path.exists(sciezka_folderu): return None
    try:
        pliki = [f for f in os.listdir(sciezka_folderu) if f.endswith('.txt')]
        if not pliki: return None
        najwiekszy_plik = max(pliki, key=lambda p: os.path.getsize(os.path.join(sciezka_folderu, p)))
        with open(os.path.join(sciezka_folderu, najwiekszy_plik), 'r', encoding='utf-8', errors='ignore') as f:
            tekst = f.read(350000) 
            return tekst
    except Exception:
        return None

def zapytaj_ai(tekst_planu, symbol, zagadnienie):
    if not tekst_planu: return "Błąd: Brak tekstu uchwały na dysku."
    
    # ==========================================
    # KROK 1: EKSTRAKTOR (Bezduszny Skaner)
    # ==========================================
    prompt_ekstrakcja = f"""
    Jesteś precyzyjnym asystentem prawnym. Twoim JEDYNYM zadaniem jest odszukanie i skopiowanie fragmentów tekstu uchwały, które mogą dotyczyć zapytania.

    ZAGADNIENIE: "{zagadnienie}"
    SYMBOL TERENU: {symbol}

    ZASADY KRYTYCZNE KROKU 1:
    1. ZERO STRESZCZANIA: Kopiuj fragmenty tekstu słowo w słowo.
    2. SZUKAJ INTELIGENTNIE: Szukaj ustaleń dla terenu {symbol} oraz ogólnych.
    3. ZWRÓĆ UWAGĘ NA LICZBY: Szukaj w szczególności fragmentów, które zawierają twarde dane (metry, m2, sztuki, odległości) dotyczące zapytania.
    4. ODRZUĆ SZUM: Omijaj zasady dla sieci infrastruktury technicznej (rury, kominy itp.).
    5. Jeśli w całym tekście nie ma NIC na ten temat, zwróć: "BRAK DANYCH".

    TEKST UCHWAŁY MPZP:
    {tekst_planu}
    """
    
    try:
        odpowiedz_skaner = client.models.generate_content(
            model=MODEL_ID, 
            contents=prompt_ekstrakcja,
            config=genai.types.GenerateContentConfig(temperature=0.0)
        )
        wyciag_z_planu = odpowiedz_skaner.text
        
        if "BRAK DANYCH" in wyciag_z_planu.upper() and len(wyciag_z_planu) < 50:
            return "W tekście tego planu nie znaleziono żadnych ustaleń dotyczących wskazanego zagadnienia."

        # ==========================================
        # KROK 2: REDAKTOR (Sztywny Szablon)
        # ==========================================
        prompt_synteza = f"""
        Otrzymujesz wyselekcjonowany wyciąg z uchwały. Masz BEZWZGLĘDNY NAKAZ zredagowania go zgodnie z poniższym szablonem.
        
        ZAGADNIENIE: "{zagadnienie}"
        SYMBOL TERENU: {symbol}
        
        ZASADY KRYTYCZNE:
        1. ZAKAZ OGÓLNIKÓW: Nie wolno Ci pisać, że "plan określa zasady". Musisz wypisać dokładnie JAKIE to są zasady, podając konkretne liczby, metry, m2 i odległości.
        2. MUSISZ użyć poniższego formatowania tekstu.
        
        WYSELEKCJONOWANY WYCIĄG Z PLANU:
        {wyciag_z_planu}
        
        --- WYMAGANY SZABLON ODPOWIEDZI (Skopiuj go i wypełnij) ---
        
        **📌 TWARDE PARAMETRY (LISTA):**
        * [Tutaj wpisz pierwszy konkretny parametr z liczbami, zakaz lub nakaz] (zgodnie z § [numer])
        * [Tutaj wpisz kolejny parametr] (zgodnie z § [numer])
        * [Jeśli w tekście są metry kwadratowe lub wysokości, MUSZĄ być w tych punktach]
        
        **📝 RZECZOWE STRESZCZENIE:**
        [Tutaj napisz krótkie, inżynieryjne podsumowanie zebranych wyżej danych w maksymalnie 3 zdaniach. Zakaz używania pustych słów urzędniczych.]
        """
        
        odpowiedz_redaktor = client.models.generate_content(
            model=MODEL_ID, 
            contents=prompt_synteza,
            # Temperatura lekko podniesiona na 0.2, żeby miał elastyczność we wpisaniu danych w szablon
            config=genai.types.GenerateContentConfig(temperature=0.2)
        )
        return odpowiedz_redaktor.text

    except Exception as e:
        return f"Błąd API: {e}"

def formatuj_tekst_planu(tekst):
    if not tekst: return "Brak tekstu."
    tekst = tekst.replace('\n', '\n\n')
    tekst = re.sub(r'(§\s*\d+[a-z]*\.?)', r'**\1**', tekst)
    tekst = re.sub(r'(?i)(rozdział\s+[A-Z0-9IVX]+|dział\s+[A-Z0-9IVX]+)', r'**\1**', tekst)
    return tekst
    
def normalizuj(tekst):
    if not tekst: return ""
    tekst = tekst.lower()
    tabela = str.maketrans("ąćęłńóśźż", "acelnoszz")
    tekst = tekst.translate(tabela)
    tekst = re.sub(r'[^a-z0-9\s]', ' ', tekst)
    return " ".join(tekst.split())

def znajdz_w_indeksie(nazwa_z_wms):
    if not nazwa_z_wms or nazwa_z_wms == "Brak planu": return None
    nazwa_szukana = normalizuj(nazwa_z_wms)
    
    for wpis in indeks_planow:
        nazwa_w_bazie = normalizuj(wpis.get('nazwa_z_mapy', ''))
        if nazwa_w_bazie == nazwa_szukana:
            return wpis.get('sciezka_folderu')
            
    slownik_nazw = {normalizuj(w.get('nazwa_z_mapy', '')): w for w in indeks_planow}
    dopasowania = difflib.get_close_matches(nazwa_szukana, slownik_nazw.keys(), n=1, cutoff=0.75)
    
    if dopasowania:
        najlepsza_nazwa = dopasowania[0]
        return slownik_nazw[najlepsza_nazwa].get('sciezka_folderu')
    return None

# --- 1. KONFIGURACJA KML I WMS ---
try: 
    fiona.drvsupport.supported_drivers['KML'] = 'rw'
    fiona.drvsupport.supported_drivers['LIBKML'] = 'rw'
except: pass

def wczytaj_kml_hybrydowo(file_obj):
    try:
        file_obj.seek(0)
        gdf = gpd.read_file(file_obj)
        if not gdf.empty: return gdf
    except: pass
    try:
        file_obj.seek(0)
        temp_path = "temp_load.kml"
        with open(temp_path, "wb") as f: f.write(file_obj.read())
        layers = fiona.listlayers(temp_path)
        all_layers = [gpd.read_file(temp_path, layer=l) for l in layers]
        if all_layers: return pd.concat(all_layers, ignore_index=True)
    except: pass
    finally:
        if os.path.exists("temp_load.kml"): os.remove("temp_load.kml")
    return gpd.GeoDataFrame()

def pobierz_wms(lat, lon):
    url = "https://wms.um.warszawa.pl/serwis"
    params = {
        "SERVICE": "WMS", "VERSION": "1.1.1", "REQUEST": "GetFeatureInfo",
        "LAYERS": "MPZP_PRZEZNACZENIE_TERENU", "QUERY_LAYERS": "MPZP_PRZEZNACZENIE_TERENU",
        "BBOX": f"{lon-0.0001},{lat-0.0001},{lon+0.0001},{lat+0.0001}",
        "SRS": "EPSG:4326", "WIDTH": "101", "HEIGHT": "101", "X": "50", "Y": "50", "INFO_FORMAT": "text/xml"
    }
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        res = requests.get(url, params=params, headers=headers, timeout=10)
        if res.status_code != 200: return None
        root = ET.fromstring(res.content)
        pelne_dane = {}
        for row in root.findall(".//ROW"):
            for child in row:
                pelne_dane[child.tag] = child.text
        if pelne_dane:
            return {
                "plan": pelne_dane.get("NAZWA_PLAN", "Brak planu"), 
                "symbol": pelne_dane.get("FUN_SYMB", "Brak symbolu"),
                "pelne_dane": pelne_dane 
            }
        return None
    except Exception:
        return None

# --- 2. INICJALIZACJA UI & PAMIĘCI ---
st.markdown("<style>header {visibility: hidden;} .block-container {padding-top: 1rem;}</style>", unsafe_allow_html=True)
if 'pts' not in st.session_state: st.session_state.pts = pd.DataFrame(columns=['lat', 'lon', 'Name'])
if 'active_idx' not in st.session_state: st.session_state.active_idx = 0
if 'batch_res' not in st.session_state: st.session_state.batch_res = None
if 'stored_results' not in st.session_state: st.session_state.stored_results = {}
if 'ai_results' not in st.session_state: st.session_state.ai_results = {}
if 'source_texts' not in st.session_state: st.session_state.source_texts = {} # <--- NOWA BAZA TEKSTÓW
if 'last_file' not in st.session_state: st.session_state.last_file = None

# --- 3. PANEL KONTROLNY ---
st.write("### 📂 1. Dane i Mapa")
c1, c2, c3 = st.columns([1.5, 1, 1.5])
with c1:
    uploaded_file = st.file_uploader("Wgraj plik KML / GeoJSON", type=['kml', 'geojson'])
    if uploaded_file:
        file_id = f"{uploaded_file.name}_{uploaded_file.size}"
        if file_id != st.session_state.last_file:
            with st.spinner("Ładuję plik..."):
                gdf = wczytaj_kml_hybrydowo(uploaded_file)
                if not gdf.empty:
                    if gdf.crs != "EPSG:4326": gdf = gdf.to_crs(epsg=4326)
                    new_df = gdf.copy()
                    new_df['lat'] = gdf.geometry.centroid.y
                    new_df['lon'] = gdf.geometry.centroid.x
                    if 'Name' not in new_df.columns: new_df['Name'] = "Punkt KML"
                    st.session_state.pts = new_df[['lat', 'lon', 'Name']]
                    st.session_state.last_file = file_id
                    st.session_state.active_idx = 0
                    st.session_state.batch_res = None
                    st.session_state.stored_results = {}
                    st.session_state.ai_results = {}
                    st.session_state.source_texts = {} # Czyszczenie pamięci tekstów przy nowym pliku
                    st.rerun()
with c2:
    view_mode = st.radio("Tryb Mapy:", ["Widok wszystkich", "Zoom na punkt"], horizontal=True)
    map_style = st.selectbox("Podkład:", ["Satelitarna", "Standardowa"])
with c3:
    st.write("🔧 Opcje AI")
    pytanie_ai = st.text_input("Pytanie do analizy:", value="Jakie są ustalenia dotyczące zabudowy mieszkaniowej i zasad lokalizacji reklam?")
    run_batch = st.button("🚀 PEŁNA ANALIZA ZBIORCZA (WMS + AI)", use_container_width=True, type="primary")

st.divider()
df = st.session_state.pts

# --- 4. MAPA ---
col_map, col_info = st.columns([1.7, 1])

def stworz_bazowa_mape(lat, lon, zoom):
    m_base = folium.Map(location=[lat, lon], zoom_start=zoom)
    Geocoder(position='topleft', add_marker=True).add_to(m_base)
    if map_style == "Satelitarna":
        folium.TileLayer('https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}', attr='Esri', name='Satelita', control=False).add_to(m_base)
    folium.WmsTileLayer(url="https://wms.um.warszawa.pl/serwis", layers="MPZP_ZAKRESY_OBOWIAZUJACE", fmt="image/png", transparent=True, name="✅ Zasięg Planów", overlay=True).add_to(m_base)
    folium.WmsTileLayer(url="https://wms.um.warszawa.pl/serwis", layers="MPZP_PRZEZNACZENIE_TERENU", fmt="image/png", transparent=True, name="🎨 Przeznaczenie Terenu", overlay=True).add_to(m_base)
    return m_base

def dodaj_warstwe_pokrycia(m):
    try:
        nazwy_baza = [normalizuj(p['nazwa_z_mapy']) for p in indeks_planow if p.get('sciezka_folderu') != 'brak']
        smieci_urzedowe = ["obszar ", "obszaru ", "rejon ", "rejonu ", "rej ", "zmiana ", "zm "]

        if os.path.exists("granice_mpzp_warszawa.geojson"):
            with open("granice_mpzp_warszawa.geojson", "r", encoding="utf-8") as f:
                dane_geojson = json.load(f)
                
            def styl_poligonu(feature):
                czy_mamy = False
                for val in feature['properties'].values():
                    if isinstance(val, str) and len(val.strip()) > 2:
                        val_znorm = normalizuj(val)
                        for nazwa_db in nazwy_baza:
                            if val_znorm == nazwa_db:
                                czy_mamy = True
                                break
                            czysta_db = nazwa_db
                            for smiec in smieci_urzedowe:
                                if czysta_db.startswith(smiec):
                                    czysta_db = czysta_db.replace(smiec, "", 1).strip()
                            if val_znorm == czysta_db:
                                czy_mamy = True
                                break
                    if czy_mamy: break
                            
                return {
                    'fillColor': '#2ECC71' if czy_mamy else '#95A5A6', 
                    'color': '#2C3E50', 'weight': 1.5, 'fillOpacity': 0.6 if czy_mamy else 0.2
                }
            
            folium.GeoJson(
                data=dane_geojson, 
                name="📊 Pokrycie Bazy (Zielone = Mamy)", 
                style_function=styl_poligonu,
                tooltip=folium.features.GeoJsonTooltip(
                    fields=['NAZWA_PLANU'], 
                    aliases=['Nazwa z mapy obowiąz. : '],
                    style="background-color: white; color: #333; font-family: sans-serif; font-size: 13px; padding: 5px; border-radius: 4px;"
                )
            ).add_to(m)
    except Exception as e:
        st.warning(f"Nie udało się załadować granic MPZP: {e}")

with col_map:
    map_data = None 
    if not df.empty:
        if st.session_state.active_idx >= len(df): st.session_state.active_idx = 0
        sel = st.selectbox("📍 Wybierz lokalizację z listy:", range(len(df)), index=st.session_state.active_idx, format_func=lambda x: f"{x+1}. {df.iloc[x]['Name']}")
        if sel != st.session_state.active_idx:
            st.session_state.active_idx = sel
            st.rerun()

        active_point = df.iloc[st.session_state.active_idx]
        m = stworz_bazowa_mape(active_point['lat'], active_point['lon'], 18)

        points_to_fit = []
        for i, row in df.iterrows():
            is_active = (i == st.session_state.active_idx)
            folium.CircleMarker([row['lat'], row['lon']], radius=6 if is_active else 4, color='red' if is_active else 'blue', fill=True, fill_opacity=0.8, tooltip=f"{i+1}. {row['Name']}").add_to(m)
            points_to_fit.append([row['lat'], row['lon']])
        if view_mode == "Widok wszystkich" and len(points_to_fit) > 1: m.fit_bounds(points_to_fit)

        dodaj_warstwe_pokrycia(m)
        folium.LayerControl().add_to(m)
        map_data = st_folium(m, use_container_width=True, height=500, key=f"map_{st.session_state.active_idx}_{view_mode}", returned_objects=["last_clicked"])
    else:
        m = stworz_bazowa_mape(52.2317, 21.0061, 11)
        dodaj_warstwe_pokrycia(m)
        folium.LayerControl().add_to(m)
        map_data = st_folium(m, use_container_width=True, height=500, key="empty_map", returned_objects=["last_clicked"])

# --- 5. PANEL INFORMACYJNY (Z MODUŁEM AI) ---
with col_info:
    if not df.empty:
        active_point = df.iloc[st.session_state.active_idx]
        current_lat, current_lon = active_point['lat'], active_point['lon']
        st.write("### 🔍 Szczegóły Punktu")
        st.write(f"**Wybrany:** {st.session_state.active_idx + 1}. {active_point['Name']}")
        
        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            if st.button("Sprawdź WMS ⚡", use_container_width=True):
                with st.spinner("Pobieram z urzędu..."):
                    res = pobierz_wms(current_lat, current_lon)
                    if res: st.session_state.stored_results[st.session_state.active_idx] = res

        res = st.session_state.stored_results.get(st.session_state.active_idx)
        
        with col_btn2:
            if res and st.button("🤖 Czytaj Plan", use_container_width=True):
                with st.spinner("Czyta się uchwałę..."):
                    
                    sciezka = znajdz_w_indeksie(res['plan'])
                    
                    if not sciezka:
                        gdf_granice = wczytaj_granice_geojson()
                        if gdf_granice is not None:
                            try:
                                pkt = Point(current_lon, current_lat)
                                znalezione = gdf_granice[gdf_granice.geometry.contains(pkt)]
                                if not znalezione.empty:
                                    for col in znalezione.columns:
                                        val = znalezione.iloc[0][col]
                                        if isinstance(val, str) and len(val) > 5:
                                            sciezka_awaryjna = znajdz_w_indeksie(val)
                                            if sciezka_awaryjna:
                                                sciezka = sciezka_awaryjna
                                                res['plan'] = val
                                                st.toast("✅ WMS kłamał, ale awaryjny GeoJSON uratował analizę!")
                                                break
                            except Exception:
                                pass
                    
                    if not sciezka:
                        st.session_state.ai_results[st.session_state.active_idx] = f"❌ BRAK W INDEKSIE: Urząd zwrócił nazwę '{res['plan']}', ale nasz słownik nie wie, gdzie leży tekst."
                    else:
                        tekst = czytaj_tekst_planu(sciezka)
                        if not tekst:
                            st.session_state.ai_results[st.session_state.active_idx] = f"📂 PUSTY FOLDER: Znalazłem ścieżkę ({sciezka}), ale brakuje tam pliku .txt z tekstem uchwały."
                        else:
                            # --- NOWOŚĆ: Zapisujemy tekst do pamięci sesji! ---
                            st.session_state.source_texts[st.session_state.active_idx] = tekst 
                            
                            nr_uchwaly = "Brak w indeksie"
                            nazwa_planu = res['plan']
                            for p in indeks_planow:
                                if p.get('sciezka_folderu') == sciezka:
                                    nr_uchwaly = p.get('nr_uchwaly', 'Brak danych')
                                    nazwa_planu = p.get('nazwa_z_mapy', res['plan'])
                                    break
                            
                            odpowiedz_ai = zapytaj_ai(tekst, res['symbol'], pytanie_ai)
                            naglowek = f"**📌 ANALIZOWANY DOKUMENT:**\n* **Plan:** {nazwa_planu}\n* **Nr Uchwały:** {nr_uchwaly}\n\n---\n"
                            
                            st.session_state.ai_results[st.session_state.active_idx] = naglowek + odpowiedz_ai

        # --- NOWOŚĆ: SYSTEM ZAKŁADEK W PANELU AI ---
        if res:
            st.success(f"**SYMBOL: {res['symbol']}**")
            st.info(f"**PLAN:** {res['plan']}")
            wynik_ai = st.session_state.ai_results.get(st.session_state.active_idx)
            
            if wynik_ai:
                # Rozdzielamy widok na Zakładki
                tab_raport, tab_zrodla = st.tabs(["📄 RAPORT AI", "📖 TEKST ŹRÓDŁOWY"])
                
                with tab_raport:
                    st.write(wynik_ai)
                    
                    # Logika pobierania Word
                    nr_u = "Brak w indeksie"
                    nazwa_p = res['plan']
                    sciezka = znajdz_w_indeksie(res['plan'])
                    if sciezka:
                        for p in indeks_planow:
                            if p.get('sciezka_folderu') == sciezka:
                                nr_u = p.get('nr_uchwaly', 'Brak danych')
                                nazwa_p = p.get('nazwa_z_mapy', res['plan'])
                                break
                    docx_file = generuj_word(wynik_ai, nr_u, nazwa_p)
                    st.download_button(
                        label="📥 Pobierz Raport (Word)",
                        data=docx_file,
                        file_name=f"Raport_{st.session_state.active_idx + 1}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                with tab_zrodla:
                    tekst_do_weryfikacji = st.session_state.source_texts.get(st.session_state.active_idx)
                    if tekst_do_weryfikacji:
                        st.caption("Przewiń poniższe okienko, aby odszukać paragrafy i zweryfikować pracę sztucznej inteligencji.")
                        # Kontener o stałej wysokości zachowuje porządek w layoucie
                        with st.container(height=450):
                            st.markdown(formatuj_tekst_planu(tekst_do_weryfikacji))
                    else:
                        st.warning("Tekst źródłowy nie jest jeszcze w pamięci. Kliknij 'Czytaj Plan', aby go załadować.")
                    
            if 'pelne_dane' in res:
                with st.expander("🕵️ Zobacz dane WMS (Detektyw)"): st.json(res['pelne_dane'])
        else:
            st.write("Oczekiwanie na pobranie danych WMS...")
    else:
        st.info("👈 Wgraj plik KML w panelu bocznym lub dodaj punkt ręcznie poniżej, aby rozpocząć analizę na mapie.")

    st.divider()
    st.write("### ➕ Dodaj nowy punkt")
    def_lat, def_lon = 52.2317, 21.0061
    if map_data and map_data.get('last_clicked'):
        def_lat, def_lon = round(map_data['last_clicked']['lat'], 6), round(map_data['last_clicked']['lng'], 6)
    c_lat, c_lon = st.columns(2)
    with c_lat: in_lat = st.text_input("Szerokość (Lat)", str(def_lat))
    with c_lon: in_lon = st.text_input("Długość (Lon)", str(def_lon))
    in_name = st.text_input("Nazwa punktu", "Ręczny Punkt")
    if st.button("Dodaj do listy 📌", use_container_width=True):
        try:
            new_row = pd.DataFrame([{'lat': float(in_lat), 'lon': float(in_lon), 'Name': in_name}])
            st.session_state.pts = pd.concat([st.session_state.pts, new_row], ignore_index=True)
            st.session_state.active_idx = len(st.session_state.pts) - 1
            st.rerun()
        except Exception:
            st.error("Błąd formatu współrzędnych!")

# --- 6. TABELA ZBIORCZA Z PEŁNĄ ANALIZĄ AI ---
if not df.empty:
    st.divider()
    if run_batch:
        st.write("### ⏳ ROZPOCZĘTO PROCEDURĘ ZBIORCZĄ")
        st.warning("Ta operacja zajmie trochę czasu (ok. 4 sekundy na punkt). Nie zamykaj okna przeglądarki.")
        bar = st.progress(0)
        status_text = st.empty()
        batch_data = []
        for i, row in df.iterrows():
            status_text.text(f"Przetwarzanie {i+1}/{len(df)}: {row['Name']}...")
            odp_wms = pobierz_wms(row['lat'], row['lon'])
            
            sciezka = None
            if odp_wms and odp_wms['plan'] != "Brak planu":
                sciezka = znajdz_w_indeksie(odp_wms['plan'])
                
                if not sciezka:
                    gdf_granice = wczytaj_granice_geojson()
                    if gdf_granice is not None:
                        try:
                            pkt = Point(row['lon'], row['lat'])
                            znalezione = gdf_granice[gdf_granice.geometry.contains(pkt)]
                            if not znalezione.empty:
                                for col in znalezione.columns:
                                    val = znalezione.iloc[0][col]
                                    if isinstance(val, str) and len(val) > 5:
                                        sciezka_awaryjna = znajdz_w_indeksie(val)
                                        if sciezka_awaryjna:
                                            sciezka = sciezka_awaryjna
                                            odp_wms['plan'] = val
                                            break
                        except Exception:
                            pass

            wynik_ai = "-"
            if not sciezka:
                wynik_ai = "❌ Brak w indeksie JSON (awaria serwerów miejskich)" if (odp_wms and odp_wms['plan'] != "Brak planu") else "Brak planu w WMS"
            else:
                tekst = czytaj_tekst_planu(sciezka)
                if not tekst:
                    wynik_ai = "📂 Brak pliku .txt na dysku"
                else:
                    # --- NOWOŚĆ: Zapisujemy tekst również podczas analizy batchowej! ---
                    st.session_state.source_texts[i] = tekst 
                    
                    sym = odp_wms['symbol'] if odp_wms else "Cały obszar"
                    wynik_ai = zapytaj_ai(tekst, sym, pytanie_ai)
                    time.sleep(4.5) 

            dane_row = {"symbol": odp_wms['symbol'] if odp_wms else "-", "plan": odp_wms['plan'] if odp_wms else "Brak danych", "Raport AI": wynik_ai}
            batch_data.append(dane_row)
            st.session_state.stored_results[i] = odp_wms if odp_wms else None
            st.session_state.ai_results[i] = wynik_ai
            bar.progress((i + 1) / len(df))
            
        status_text.text("✅ Analiza zakończona sukcesem!")
        st.session_state.batch_res = pd.DataFrame(batch_data)
        st.rerun()

    df_tab = df.copy()
    if st.session_state.batch_res is not None:
        res_cols = st.session_state.batch_res.rename(columns={'plan':'NAZWA PLANU', 'symbol':'SYMBOL'})
        if 'pelne_dane' in res_cols.columns: res_cols = res_cols.drop(columns=['pelne_dane'])
        df_tab = pd.concat([res_cols, df_tab.reset_index(drop=True)], axis=1)
    else:
        df_tab['NAZWA PLANU'] = "-"; df_tab['SYMBOL'] = "-"; df_tab['Raport AI'] = "-"
    
    df_tab.insert(0, 'Lp.', range(1, len(df_tab) + 1))
    
    st.write("### 📊 Wyniki Analizy")
    
    if st.session_state.batch_res is not None:
        col_dl1, col_dl2, col_dl3 = st.columns(3)
        
        df_do_eksportu = df_tab[['Lp.', 'SYMBOL', 'NAZWA PLANU', 'Name', 'Raport AI']]
        
        with col_dl1:
            st.download_button(
                label="📝 Pobierz Raport (Word)",
                data=generuj_word_zbiorczy(df_do_eksportu),
                file_name=f"Zbiorczy_Raport_MPZP_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        with col_dl2:
            st.download_button(
                label="📊 Pobierz Zestawienie (Excel)",
                data=eksport_do_excel(df_do_eksportu),
                file_name=f"Zestawienie_MPZP_{datetime.now().strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        with col_dl3:
            st.download_button(
                label="📄 Pobierz Zestawienie (CSV)",
                data=eksport_do_csv(df_do_eksportu),
                file_name=f"Zestawienie_MPZP_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv",
                use_container_width=True
            )
            
    event = st.dataframe(df_tab[['Lp.', 'SYMBOL', 'NAZWA PLANU', 'Name', 'Raport AI']], on_select="rerun", selection_mode="single-row", hide_index=True, use_container_width=True)
    if event and event.selection and event.selection.rows:
        row = event.selection.rows[0]
        if row != st.session_state.active_idx:
            st.session_state.active_idx = row
            st.rerun()